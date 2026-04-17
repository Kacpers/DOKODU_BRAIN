#!/usr/bin/env python3
"""
Dokodu Offer Generator
Generuje brandowany .docx z pliku Markdown oferty.

Użycie:
    python3 generate_offer.py <input.md> [output.docx]

Przykład:
    python3 generate_offer.py ../20_AREAS/AREA_Customers/_PROSPECTS/Korollo_Oferta_2026-03.md
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Kolory Dokodu ──────────────────────────────────────────────
NAVY     = RGBColor(0x0F, 0x21, 0x37)   # #0F2137 — główny
RED      = RGBColor(0xE6, 0x39, 0x46)   # #E63946 — akcent
SLATE    = RGBColor(0x5A, 0x66, 0x77)   # #5A6677 — secondary text
LIGHT_BG = RGBColor(0xF8, 0xFA, 0xFC)   # #F8FAFC — tło sekcji
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
BORDER   = RGBColor(0xE2, 0xE8, 0xF0)   # #E2E8F0 — linie

FONT_NAME = "Lato"


# ── Helpery XML ────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), kwargs.get(edge, "none"))
        border.set(qn("w:sz"), kwargs.get("sz", "4"))
        border.set(qn("w:color"), kwargs.get("color", "E2E8F0"))
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_paragraph_shading(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def set_para_border_left(para, color="E63946", sz="24"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


def set_page_margins(doc, top=2.5, bottom=2.5, left=2.5, right=2.5):
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


def add_run(para, text, bold=False, italic=False, size=11,
            color=None, font=FONT_NAME):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name  = font
    run.font.size  = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


# ── Sekcje dokumentu ──────────────────────────────────────────
def add_header_band(doc, client_name: str, subtitle: str):
    """Granatowy pasek nagłówkowy z nazwą klienta."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    left_cell  = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    set_cell_bg(left_cell,  "0F2137")
    set_cell_bg(right_cell, "0F2137")

    # Usuń obramowanie tabeli
    for cell in [left_cell, right_cell]:
        set_cell_border(cell, top="none", left="none", bottom="none", right="none")

    # Lewa strona — nazwa propozycji
    lp = left_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_before = Pt(14)
    lp.paragraph_format.space_after  = Pt(4)
    add_run(lp, "Propozycja Współpracy", bold=False, size=10, color=WHITE)

    lp2 = left_cell.add_paragraph()
    lp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp2.paragraph_format.space_before = Pt(0)
    lp2.paragraph_format.space_after  = Pt(14)
    add_run(lp2, client_name, bold=True, size=22, color=RED)

    lp3 = left_cell.add_paragraph()
    lp3.paragraph_format.space_before = Pt(0)
    lp3.paragraph_format.space_after  = Pt(12)
    add_run(lp3, subtitle, bold=False, size=10, color=RGBColor(0xA0, 0xAE, 0xBE))

    # Prawa strona — Dokodu
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(20)
    rp.paragraph_format.space_after  = Pt(4)
    add_run(rp, "dokodu", bold=True, size=20, color=WHITE)
    add_run(rp, ".it", bold=True, size=20, color=RED)

    rp2 = right_cell.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp2.paragraph_format.space_before = Pt(0)
    rp2.paragraph_format.space_after  = Pt(14)
    add_run(rp2, "kacper@dokodu.it  ·  +48 508 106 046", size=9,
            color=RGBColor(0xA0, 0xAE, 0xBE))

    doc.add_paragraph()  # odstęp


def add_section_heading(doc, text: str, level=1):
    """Nagłówek sekcji."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18) if level == 1 else Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    if level == 1:
        add_run(p, text.upper(), bold=True, size=11, color=NAVY)
        # Linia pod nagłówkiem
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "E63946")
        pBdr.append(bottom)
        pPr.append(pBdr)
    else:
        add_run(p, text, bold=True, size=11, color=NAVY)


def add_body_text(doc, text: str, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    # Obsługa **bold** w tekście
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        if part:
            add_run(p, part, bold=(i % 2 == 1), size=10,
                    color=NAVY if (i % 2 == 1) else RGBColor(0x1A, 0x2B, 0x3C))
    return p


def add_bullet(doc, text: str, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.4 + level * 0.5)
    parts = re.split(r'\*\*(.+?)\*\*', text.lstrip("- •·[ ]✓☐"))
    for i, part in enumerate(parts):
        if part:
            add_run(p, part, bold=(i % 2 == 1), size=10)


def add_quote_block(doc, text: str):
    """Blok cytatu z czerwoną pionową linią."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    add_paragraph_shading(p, "F8FAFC")
    set_para_border_left(p)
    add_run(p, text.lstrip("> ").strip(), italic=True, size=10,
            color=RGBColor(0x2D, 0x3E, 0x50))


def add_timeline_block(doc, lines: list):
    """Blok harmonogramu — tabela 2-kolumnowa."""
    if not lines:
        return
    table = doc.add_table(rows=len(lines), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    col_widths = [Cm(4), Cm(12)]
    for i, row in enumerate(table.rows):
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]
        bg = "F0F4F8" if i % 2 == 0 else "FFFFFF"
        set_cell_bg(row.cells[0], "0F2137" if i == 0 else bg)
        set_cell_bg(row.cells[1], "0F2137" if i == 0 else bg)

        for j, cell in enumerate(row.cells):
            set_cell_border(cell, top="single", left="none",
                            bottom="single", right="none",
                            color="E2E8F0", sz="4")
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_before = Pt(4)
            cp.paragraph_format.space_after  = Pt(4)
            if j < len(lines[i]):
                c = WHITE if i == 0 else NAVY
                add_run(cp, lines[i][j], bold=(i == 0), size=9, color=c)

    doc.add_paragraph()


def add_pricing_table(doc, rows: list):
    """Tabela cenowa z granatowym nagłówkiem."""
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        is_header = i == 0
        is_total  = "Łącznie" in str(row_data[0]) or "**" in str(row_data[0])

        bg_left  = "0F2137" if is_header else ("EBF0F5" if is_total else "FFFFFF")
        bg_right = "0F2137" if is_header else ("EBF0F5" if is_total else "FFFFFF")

        set_cell_bg(row.cells[0], bg_left)
        set_cell_bg(row.cells[1], bg_right)

        for j, cell in enumerate(row.cells):
            set_cell_border(cell, top="single", left="none",
                            bottom="single", right="none",
                            color="E2E8F0", sz="4")
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_before = Pt(5)
            cp.paragraph_format.space_after  = Pt(5)
            cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j == 1 else WD_ALIGN_PARAGRAPH.LEFT
            text = re.sub(r'\*\*|\*', '', str(row_data[j]) if j < len(row_data) else "")
            color = WHITE if is_header else (NAVY if is_total else RGBColor(0x2D, 0x3E, 0x50))
            add_run(cp, text, bold=(is_header or is_total), size=10, color=color)

    doc.add_paragraph()


def add_footer(doc):
    """Stopka z danymi Dokodu."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_paragraph_shading(p, "0F2137")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.3)
    add_run(p, "Dokodu  ·  ", bold=True, size=9, color=WHITE)
    add_run(p, "kacper@dokodu.it  ·  +48 508 106 046  ·  dokodu.it  ·  ul. Kosynierów 76/22, 84-230 Rumia",
            size=9, color=RGBColor(0xA0, 0xAE, 0xBE))


# ── Parser Markdown → sekcje ──────────────────────────────────
def parse_and_render(doc, md_text: str):
    lines = md_text.split("\n")
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # Pomiń frontmatter YAML
        if line.strip() == "---" and i == 0:
            i += 1
            while i < len(lines) and lines[i].strip() != "---":
                i += 1
            i += 1
            continue

        # Poziomy poziome ---
        if re.match(r'^---+$', line.strip()):
            i += 1
            continue

        # Blok kodu (harmonogram)
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                # Renderuj jako timeline
                tl_rows = []
                for cl in code_lines:
                    if "│" in cl:
                        parts = cl.split("│", 1)
                        tl_rows.append([parts[0].strip(), parts[1].strip()])
                    elif cl.strip():
                        tl_rows.append([cl.strip(), ""])
                if tl_rows:
                    add_timeline_block(doc, tl_rows)
                in_code_block = False
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Nagłówki
        if line.startswith("# ") and not line.startswith("## "):
            # H1 — tylko tytuł główny, już w headerze
            i += 1
            continue

        if line.startswith("## "):
            add_section_heading(doc, line[3:].strip(), level=1)
            i += 1
            continue

        if line.startswith("### "):
            add_section_heading(doc, line[4:].strip(), level=2)
            i += 1
            continue

        if line.startswith("#### "):
            add_section_heading(doc, line[5:].strip(), level=2)
            i += 1
            continue

        # Cytaty
        if line.startswith("> "):
            add_quote_block(doc, line)
            i += 1
            continue

        # Listy z checkboxem [ ] lub [x]
        if re.match(r'^- \[[ x]\]', line):
            text = re.sub(r'^- \[[ x]\] ', '', line)
            add_bullet(doc, text)
            i += 1
            continue

        # Listy zwykłe
        if re.match(r'^[-*] ', line) or re.match(r'^\d+\. ', line):
            text = re.sub(r'^[-*] |^\d+\. ', '', line)
            add_bullet(doc, text)
            i += 1
            continue

        # Tabele Markdown
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                if not re.match(r'^\|[-| :]+\|$', lines[i]):
                    table_lines.append(lines[i])
                i += 1
            # Sprawdź czy to tabela cenowa (zawiera PLN)
            is_pricing = any("PLN" in l or "netto" in l.lower() for l in table_lines)
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip("|").split("|")]
                rows.append(cells)
            if is_pricing:
                add_pricing_table(doc, rows)
            else:
                add_pricing_table(doc, rows)  # używamy tego samego stylu
            continue

        # Pusta linia
        if not line.strip():
            i += 1
            continue

        # Pogrubiony tekst jako emphasis (bold + indent)
        if line.startswith("**") and line.endswith("**") and len(line) > 4:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            add_run(p, line.strip("*"), bold=True, size=10, color=NAVY)
            i += 1
            continue

        # Tekst blokowy > (bez cytatu) — info box
        if line.startswith("*") and line.endswith("*"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.left_indent  = Cm(0.3)
            add_paragraph_shading(p, "F8FAFC")
            add_run(p, line.strip("*"), italic=True, size=9, color=SLATE)
            i += 1
            continue

        # Zwykły paragraf
        add_body_text(doc, line.strip())
        i += 1


# ── Główna funkcja ────────────────────────────────────────────
def generate_offer(input_path: str, output_path: str = None):
    md_path = Path(input_path)
    if not md_path.exists():
        print(f"❌ Plik nie istnieje: {input_path}")
        sys.exit(1)

    md_text = md_path.read_text(encoding="utf-8")

    # Wyciągnij tytuł klienta z H1 lub frontmatter
    client_match = re.search(r'^# Propozycja Współpracy: (.+)$', md_text, re.MULTILINE)
    client_name  = client_match.group(1).strip() if client_match else md_path.stem

    subtitle_match = re.search(r'^\*\*Dokodu × .+\*\*$', md_text, re.MULTILINE)
    subtitle = subtitle_match.group(0).strip("*") if subtitle_match else "Marzec 2026"

    if output_path is None:
        output_path = str(md_path.with_suffix(".docx"))

    doc = Document()
    set_page_margins(doc, top=0, bottom=2, left=0, right=0)

    # Ustaw domyślny styl
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10)

    # Header
    add_header_band(doc, client_name, subtitle)

    # Ustaw marginesy treści po headerze
    section = doc.sections[0]
    section.left_margin  = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Treść
    parse_and_render(doc, md_text)

    # Stopka
    add_footer(doc)

    doc.save(output_path)
    print(f"✅ Oferta wygenerowana: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Użycie: python3 generate_offer.py <input.md> [output.docx]")
        sys.exit(1)
    inp = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else None
    generate_offer(inp, out)
