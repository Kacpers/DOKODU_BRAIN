"""
POC: AI-powered uzupełnianie dokumentów przetargowych.
Wysyła treść dokumentu do Claude, dostaje z powrotem mapowanie pól.
"""
import os
import json
import anthropic
from docx import Document

COMPANY_PROFILE = {
    "nazwa_firmy": "DOKODU Sp. z o.o.",
    "nip": "5882473305",
    "regon": "520149113",
    "krs": "0000925166",
    "adres": "ul. Kosynierów 76/22, 84-230 Rumia",
    "email": "biuro@dokodu.it",
    "telefon": "+48 508 106 046",
    "osoba_kontaktowa": "Kacper Sieradziński",
    "stanowisko": "CEO",
    "status_przedsiebiorcy": "mikroprzedsiębiorstwo",
    "kapital_zakladowy": "5 000 PLN",
    "miejscowosc": "Rumia",
    "data": "08.04.2026",
}

def extract_doc_text(doc: Document) -> str:
    """Wyciąga pełną treść dokumentu z oznaczeniami struktury."""
    parts = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            parts.append(f"[P{i}] {text}")

    for t_idx, table in enumerate(doc.tables):
        parts.append(f"\n[TABLE {t_idx}]")
        for r_idx, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            parts.append(f"  [T{t_idx}R{r_idx}] {' | '.join(cells)}")

    return "\n".join(parts)


def ask_ai_for_fields(doc_text: str, filename: str) -> list[dict]:
    """Pyta Claude o pola do wypełnienia."""
    client = anthropic.Anthropic()

    prompt = f"""Analizujesz formularz przetargowy "{filename}".

Profil firmy wykonawcy:
{json.dumps(COMPANY_PROFILE, ensure_ascii=False, indent=2)}

Treść dokumentu (z oznaczeniami pozycji):
---
{doc_text}
---

Znajdź WSZYSTKIE pola w dokumencie, które powinny być wypełnione danymi wykonawcy.
Pola to: puste komórki tabel obok etykiet, wielokropki (……), kropki (......), puste nawiasy.

Dla każdego znalezionego pola zwróć JSON:
- "location_type": "table_cell" lub "paragraph"
- "location_id": np. "T0R3" (tabela 0, wiersz 3) lub "P13" (paragraf 13)
- "label": etykieta pola (co opisuje to pole)
- "current_value": co jest teraz w polu (puste, wielokropki, itp.)
- "suggested_value": wartość z profilu firmy do wpisania
- "confidence": "high", "medium", "low"
- "reasoning": krótkie uzasadnienie

Zwróć TYLKO tablicę JSON, bez dodatkowego tekstu. Jeśli pole nie pasuje do żadnych danych z profilu (np. cena, specyfikacja techniczna), NIE dodawaj go."""

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )

    text = response.content[0].text.strip()
    # Wyciągnij JSON z odpowiedzi
    if text.startswith("```"):
        text = text.split("```")[1]
        if text.startswith("json"):
            text = text[4:]

    return json.loads(text)


def main():
    input_dir = "/mnt/d/DOWNLOAD/reprobaowycenwykonaniemoduuuzupenianiadanychw"

    results = []
    for fname in sorted(os.listdir(input_dir)):
        if not fname.endswith('.docx'):
            continue
        fpath = os.path.join(input_dir, fname)
        try:
            doc = Document(fpath)
            doc_text = extract_doc_text(doc)
            fields = ask_ai_for_fields(doc_text, fname)
            results.append({"file": fname, "fields": fields})
        except Exception as e:
            results.append({"file": fname, "error": str(e)})

    # Raport
    print("=" * 70)
    print("POC AI-POWERED — RAPORT")
    print("=" * 70)

    total_fields = 0
    for r in results:
        if "error" in r:
            print(f"\n❌ {r['file']}: {r['error']}")
            continue

        fields = r["fields"]
        total_fields += len(fields)

        print(f"\n📄 {r['file']}")
        print(f"   Znalezione pola: {len(fields)}")

        for f in fields:
            conf_icon = {"high": "🟢", "medium": "🟡", "low": "🔴"}.get(f.get("confidence", ""), "⚪")
            print(f"   {conf_icon} [{f.get('location_id', '?')}] {f.get('label', '?')}")
            print(f"      Obecna wartość: {f.get('current_value', '?')[:60]}")
            print(f"      Sugerowana: {f.get('suggested_value', '?')}")
            print(f"      Powód: {f.get('reasoning', '?')[:80]}")

    print(f"\n{'=' * 70}")
    print(f"PODSUMOWANIE: AI znalazło {total_fields} pól w {len(results)} dokumentach")
    print("=" * 70)

if __name__ == "__main__":
    main()
