"""Write filled values back into DOCX preserving formatting."""
import re
from pathlib import Path
from docx import Document
from .models import FieldResult

_PLACEHOLDER_RE = re.compile(r'[…\.]{2,}|_{3,}')


def _parse_table_location(loc_id: str) -> tuple[int, int] | None:
    m = re.match(r"T(\d+)R(\d+)", loc_id)
    return (int(m.group(1)), int(m.group(2))) if m else None


def _parse_paragraph_location(loc_id: str) -> int | None:
    m = re.match(r"P(\d+)", loc_id)
    return int(m.group(1)) if m else None


def _is_placeholder(text: str) -> bool:
    """Check if text is a placeholder (dots, ellipsis, underscores, empty)."""
    stripped = text.strip()
    if not stripped:
        return True
    return bool(re.match(r'^[\s…\._□\-]*$', stripped))


def _replace_placeholder_in_paragraph(paragraph, value: str) -> bool:
    """Replace placeholder content in paragraph with value."""
    full_text = paragraph.text

    # 1. Replace first placeholder, clear all remaining ones across runs
    if _PLACEHOLDER_RE.search(full_text):
        inserted = False
        found_placeholder_zone = False
        for run in paragraph.runs:
            if _PLACEHOLDER_RE.search(run.text):
                found_placeholder_zone = True
                if not inserted:
                    run.text = _PLACEHOLDER_RE.sub(value, run.text, count=1)
                    run.text = _PLACEHOLDER_RE.sub("", run.text)
                    inserted = True
                else:
                    run.text = _PLACEHOLDER_RE.sub("", run.text)
            elif found_placeholder_zone and re.match(r'^[\s\.…]*$', run.text):
                # Trailing fragment of placeholder zone (single dots, spaces)
                run.text = ""
        return True

    # 2. Handle parenthesized descriptions like "(nazwa lub pieczęć wykonawcy)"
    if re.search(r'^\s*\(.*\)\s*$', full_text):
        _set_paragraph_text(paragraph, value)
        return True

    # 3. Handle "Label:" pattern — append value after colon
    if re.match(r'^.+:\s*$', full_text.strip()):
        if paragraph.runs:
            paragraph.runs[-1].text = paragraph.runs[-1].text.rstrip() + " " + value
        else:
            paragraph.text = full_text.rstrip() + " " + value
        return True

    # 4. Handle fully empty paragraph
    if not full_text.strip():
        _set_paragraph_text(paragraph, value)
        return True

    return False


def _set_paragraph_text(paragraph, text: str) -> None:
    """Set paragraph text, preserving first run's formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = text


def _write_table_cell(doc: Document, field: FieldResult) -> None:
    loc = _parse_table_location(field.location_id)
    if not loc:
        return
    t_idx, r_idx = loc
    try:
        row = doc.tables[t_idx].rows[r_idx]
        cells = row.cells
        target_cell = None

        if len(cells) >= 2:
            c0_placeholder = _is_placeholder(cells[0].text)
            c1_placeholder = _is_placeholder(cells[1].text)

            if c0_placeholder and c1_placeholder:
                # Both cells are placeholders — fill cell[0] (left), it's typically the first value field
                target_cell = cells[0]
            elif c1_placeholder:
                # Standard label|value pattern
                target_cell = cells[1]
            elif not cells[1].text.strip():
                target_cell = cells[1]
            else:
                # Cell[1] has content — check if it matches original_value
                if field.original_value and cells[1].text.strip() == field.original_value.strip():
                    target_cell = cells[1]
        elif len(cells) == 1:
            if _is_placeholder(cells[0].text):
                target_cell = cells[0]

        if target_cell is None:
            # Last resort for 2+ cell rows: use cell[1]
            if len(cells) >= 2:
                target_cell = cells[1]
            else:
                return

        if target_cell.paragraphs:
            # Find the right paragraph — prefer one with placeholder, fallback to first
            filled = False
            for para in target_cell.paragraphs:
                if _is_placeholder(para.text) and para.text.strip():
                    _set_paragraph_text(para, field.filled_value)
                    filled = True
                    break
            if not filled:
                # Clear all paragraphs and set first one
                target_cell.paragraphs[0].text = field.filled_value
                for para in target_cell.paragraphs[1:]:
                    _set_paragraph_text(para, "")
        else:
            target_cell.text = field.filled_value
    except (IndexError, KeyError):
        return


def _write_paragraph(doc: Document, field: FieldResult) -> None:
    p_idx = _parse_paragraph_location(field.location_id)
    if p_idx is None:
        return
    try:
        para = doc.paragraphs[p_idx]
        _replace_placeholder_in_paragraph(para, field.filled_value)
    except IndexError:
        return


def write_fields(source_path: Path, fields: list[FieldResult], output_path: Path) -> None:
    """Write field values into DOCX and save to output_path."""
    doc = Document(str(source_path))

    for field in fields:
        if field.location_type == "table_cell":
            _write_table_cell(doc, field)
        elif field.location_type == "paragraph":
            _write_paragraph(doc, field)

    doc.save(str(output_path))
