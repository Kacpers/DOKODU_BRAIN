import pytest
from pathlib import Path
from docx import Document
from backend.docx_writer import write_fields
from backend.models import FieldResult, Confidence


def test_write_table_cell(simple_table_docx, tmp_path):
    fields = [
        FieldResult(
            location_id="T0R0", location_type="table_cell", label="Firma",
            original_value="", filled_value="Test Corp",
            confidence=Confidence.HIGH,
        ),
        FieldResult(
            location_id="T0R1", location_type="table_cell", label="NIP",
            original_value="", filled_value="1234567890",
            confidence=Confidence.HIGH,
        ),
    ]
    out_path = tmp_path / "filled.docx"
    write_fields(simple_table_docx, fields, out_path)
    doc = Document(str(out_path))
    assert doc.tables[0].rows[0].cells[1].text.strip() == "Test Corp"
    assert doc.tables[0].rows[1].cells[1].text.strip() == "1234567890"


def test_write_paragraph_dots(paragraph_dots_docx, tmp_path):
    fields = [
        FieldResult(
            location_id="P1", location_type="paragraph", label="NIP",
            original_value="……………………………………", filled_value="1234567890",
            confidence=Confidence.HIGH,
        ),
    ]
    out_path = tmp_path / "filled.docx"
    write_fields(paragraph_dots_docx, fields, out_path)
    doc = Document(str(out_path))
    assert "1234567890" in doc.paragraphs[1].text
    assert "……" not in doc.paragraphs[1].text
