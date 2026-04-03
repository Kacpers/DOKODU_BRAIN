import pytest
from pathlib import Path
from docx import Document


@pytest.fixture(autouse=True)
def isolate_profile(tmp_path, monkeypatch):
    """Redirect profile storage to tmp_path so tests don't mutate default_profile.json."""
    import backend.profile as prof
    test_profile_path = tmp_path / "test_profile.json"
    import json
    test_profile_path.write_text(json.dumps({
        "nazwa_firmy": "TestFirma Sp. z o.o.", "nip": "1111111111", "regon": "222222222",
        "krs": "0000333333", "adres": "ul. Testowa 5, 00-001 Warszawa",
        "adres_korespondencyjny": "ul. Testowa 5, 00-001 Warszawa",
        "email": "test@firma.pl", "telefon": "+48 111 222 333",
        "osoba_kontaktowa": "Anna Testowa", "stanowisko": "CEO",
        "status_przedsiebiorcy": "mikro", "miejscowosc": "Warszawa"
    }, ensure_ascii=False))
    monkeypatch.setattr(prof, "PROFILE_PATH", test_profile_path)


@pytest.fixture
def sample_profile():
    from backend.models import CompanyProfile
    return CompanyProfile(
        nazwa_firmy="TestFirma Sp. z o.o.",
        nip="1111111111",
        regon="222222222",
        krs="0000333333",
        adres="ul. Testowa 5, 00-001 Warszawa",
        adres_korespondencyjny="ul. Testowa 5, 00-001 Warszawa",
        email="test@firma.pl",
        telefon="+48 111 222 333",
        osoba_kontaktowa="Anna Testowa",
        stanowisko="CEO",
        status_przedsiebiorcy="mikro",
        miejscowosc="Warszawa",
    )


@pytest.fixture
def simple_table_docx(tmp_path) -> Path:
    """DOCX with a table containing company fields."""
    doc = Document()
    table = doc.add_table(rows=4, cols=2)
    fields = [("Firma", ""), ("NIP", ""), ("Adres", ""), ("E-mail", "")]
    for i, (label, val) in enumerate(fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = val
    path = tmp_path / "simple_table.docx"
    doc.save(path)
    return path


@pytest.fixture
def paragraph_dots_docx(tmp_path) -> Path:
    """DOCX with paragraph-style fields using dots."""
    doc = Document()
    doc.add_paragraph("FORMULARZ OFERTOWY")
    doc.add_paragraph("NIP: ……………………………………")
    doc.add_paragraph("tel.: .....................................................")
    doc.add_paragraph("e-mail: ................................................")
    doc.add_paragraph("Łącznie brutto ……………………... PLN")
    path = tmp_path / "paragraph_dots.docx"
    doc.save(path)
    return path
