"""AI engine: Claude API for field detection in tender documents."""
import json
from datetime import date
from pathlib import Path
from docx import Document
import anthropic
from ..models import CompanyProfile, FieldResult, Confidence

_TIMEOUT = 120


def _extract_doc_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    parts = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            parts.append(f"[P{i}] {text}")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            parts.append(f"[T{t_idx}R{r_idx}] {' | '.join(cells)}")
    return "\n".join(parts)


def _build_prompt(doc_text: str, profile: CompanyProfile, filename: str) -> str:
    p = profile
    today = date.today().strftime("%d.%m.%Y")
    contact_line = ", ".join(filter(None, [p.osoba_kontaktowa, p.telefon, p.email]))

    profile_data = {
        "nazwa_firmy": p.nazwa_firmy,
        "nip": p.nip,
        "regon": p.regon,
        "krs": p.krs,
        "adres": p.adres,
        "adres_korespondencyjny": p.adres_korespondencyjny or p.adres,
        "email": p.email,
        "telefon": p.telefon,
        "osoba_kontaktowa": p.osoba_kontaktowa,
        "stanowisko": p.stanowisko,
        "status_przedsiebiorcy": p.status_przedsiebiorcy,
        "miejscowosc": p.miejscowosc,
        "miejscowosc_i_data": f"{p.miejscowosc}, {today}",
        "kontakt_pelny": contact_line,
        "regon_nip_krs": f"{p.regon} / {p.nip} / {p.krs}",
        "nazwa_i_adres": f"{p.nazwa_firmy}, {p.adres}",
    }
    profile_json = json.dumps(profile_data, ensure_ascii=False, indent=2)

    return f"""Analizujesz formularz przetargowy "{filename}".

Dane firmy wykonawcy do wpisania:
{profile_json}

Treść dokumentu (z oznaczeniami pozycji):
[P0] = paragraf o indeksie 0
[T0R3] = tabela 0, wiersz 3
---
{doc_text}
---

ZADANIE: Znajdź WSZYSTKIE pola w dokumencie, które powinny być wypełnione danymi wykonawcy z profilu powyżej.

GDZIE SZUKAĆ PÓL:
1. TABELE: puste komórki obok etykiet (np. "Nazwa Wykonawcy | " — pusta druga komórka)
2. PARAGRAFY z wielokropkami: "NIP: ............" lub "e-mail: ……………"
3. PARAGRAFY z podkreśleniami: "Nazwa firmy: ________________"
4. PARAGRAFY z pustymi nawiasami lub opisami w nawiasach: "(nazwa lub pieczęć wykonawcy)"
5. TABELE ze złożonymi etykietami: "REGON / NIP/ KRS" → wpisz "520149113 / 5882473305 / 0000925166"
6. Pola podpisu: "miejscowość i data", "pieczęć wykonawcy", "pieczęć Oferenta"
7. Pola wieloliniowe: etykieta może zawierać newline, np. "Osoba do kontaktu\n(imię i nazwisko)"

CZEGO NIE WYPEŁNIAĆ:
- Pola dotyczące cen, kwot, wartości oferty
- Pola techniczne (specyfikacje, parametry)
- Checkboxy TAK/NIE
- Pola już wypełnione prawidłowymi danymi
- Pola dotyczące zamawiającego (nie wykonawcy)
- Numery pozycji, nagłówki tabel

WAŻNE:
- Dla pola "status przedsiębiorcy" z opcjami "mikro/małe/średnie/duże" → wpisz "{p.status_przedsiebiorcy}"
- Dla pól łączonych (imię + telefon + email) → użyj pełnego kontaktu: "{contact_line}"
- Dla "miejscowość i data" / "miejscowość, data" → wpisz "{p.miejscowosc}, {today}"
- Dla "nazwa lub pieczęć wykonawcy" / "pieczęć wykonawcy" → wpisz "{p.nazwa_firmy}"
- Dla "REGON / NIP/ KRS" → wpisz "{p.regon} / {p.nip} / {p.krs}"
- Jeśli pole pyta o "Nazwa i adres" razem → wpisz "{p.nazwa_firmy}, {p.adres}"

Dla każdego pola zwróć:
- location_id: dokładne oznaczenie z dokumentu (np. "T0R3" lub "P13")
- location_type: "table_cell" lub "paragraph"
- label: etykieta pola jak w dokumencie
- current_value: co jest teraz w polu (wielokropki, puste, etc.)
- suggested_value: wartość do wpisania z profilu
- confidence: "high" jeśli jednoznaczne dopasowanie, "medium" jeśli prawdopodobne, "low" jeśli niepewne

Zwróć TYLKO pola dla których masz dane w profilu. Nie zgaduj wartości."""


_EXTRACT_FIELDS_TOOL = {
    "name": "extract_fields",
    "description": "Return all detected company fields from the tender document.",
    "input_schema": {
        "type": "object",
        "properties": {
            "fields": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "location_id": {"type": "string", "description": "e.g. T0R3 or P13"},
                        "location_type": {"type": "string", "enum": ["table_cell", "paragraph"]},
                        "label": {"type": "string"},
                        "current_value": {"type": "string"},
                        "suggested_value": {"type": "string"},
                        "confidence": {"type": "string", "enum": ["high", "medium", "low"]},
                    },
                    "required": ["location_id", "location_type", "label", "suggested_value", "confidence"],
                },
            }
        },
        "required": ["fields"],
    },
}


def _call_claude(prompt: str) -> list[dict]:
    client = anthropic.Anthropic()
    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=8192,
        timeout=_TIMEOUT,
        tools=[_EXTRACT_FIELDS_TOOL],
        tool_choice={"type": "tool", "name": "extract_fields"},
        messages=[{"role": "user", "content": prompt}],
    )
    for block in response.content:
        if block.type == "tool_use" and block.name == "extract_fields":
            return block.input.get("fields", [])
    return []


def analyze_ai(docx_path: Path, profile: CompanyProfile) -> list[FieldResult]:
    """Analyze DOCX using Claude AI and return detected fields."""
    doc_text = _extract_doc_text(docx_path)
    filename = Path(docx_path).name
    prompt = _build_prompt(doc_text, profile, filename)
    raw_fields = _call_claude(prompt)

    results = []
    for f in raw_fields:
        try:
            results.append(FieldResult(
                location_id=f["location_id"],
                location_type=f.get("location_type", "paragraph"),
                label=f["label"],
                original_value=f.get("current_value", ""),
                filled_value=f["suggested_value"],
                confidence=Confidence(f.get("confidence", "medium")),
            ))
        except (KeyError, ValueError):
            continue
    return results
