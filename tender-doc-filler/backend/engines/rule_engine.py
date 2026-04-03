"""Rule-based engine: regex + table cell matching for known company fields."""
import re
from datetime import date
from pathlib import Path
from docx import Document
from ..models import CompanyProfile, FieldResult, FieldSource, Confidence


def _build_label_map(profile: CompanyProfile) -> dict[str, str]:
    """Map normalized label strings to profile values."""
    p = profile
    contact_line = ", ".join(filter(None, [
        p.osoba_kontaktowa, p.telefon, p.email,
    ]))
    entries = {
        # Nazwa firmy
        "firma": p.nazwa_firmy,
        "nazwa": p.nazwa_firmy,
        "nazwa firmy": p.nazwa_firmy,
        "nazwa wykonawcy": p.nazwa_firmy,
        "nazwa oferenta": p.nazwa_firmy,
        # Identyfikatory
        "nip": p.nip,
        "regon": p.regon,
        "krs": p.krs,
        # Adres
        "adres": p.adres,
        "adres wykonawcy": p.adres,
        "adres do korespondencji": p.adres_korespondencyjny,
        "adres korespondencyjny": p.adres_korespondencyjny,
        "siedziba i adres": p.adres,
        "siedziba": p.adres,
        "adres siedziby": p.adres,
        # Email
        "adres e-mail": p.email,
        "e-mail": p.email,
        "email": p.email,
        # Telefon
        "telefon": p.telefon,
        "tel": p.telefon,
        "numer telefonu": p.telefon,
        # Osoba kontaktowa
        "osoba do kontaktu": p.osoba_kontaktowa,
        "osoba do kontaktu\n(imię i nazwisko)": p.osoba_kontaktowa,
        "osoba kontaktowa": p.osoba_kontaktowa,
        "imię i nazwisko": p.osoba_kontaktowa,
        "imię, nazwisko, telefon i e-mail osoby do kontaktu": contact_line,
        # Forma prawna / status
        "forma prawna": p.status_przedsiebiorcy,
        "status przedsiębiorcy": p.status_przedsiebiorcy,
        "status przedsiebiorcy": p.status_przedsiebiorcy,
        "wielkość przedsiębiorstwa": p.status_przedsiebiorcy,
        # Data i miejsce
        "miejscowość i data": f"{p.miejscowosc}, {date.today().strftime('%d.%m.%Y')}",
        "miejscowość": p.miejscowosc,
        "miejscowosc": p.miejscowosc,
    }
    return {k: v for k, v in entries.items() if v}


def _normalize_label(text: str) -> str:
    t = text.strip().lower()
    t = re.sub(r'[:\*\.]+$', '', t).strip()
    t = re.sub(r'\s+', ' ', t)
    return t


def _is_empty_value(text: str) -> bool:
    stripped = text.strip()
    return (
        stripped == ""
        or bool(re.match(r'^[…\.]+$', stripped))
        or "mikro/małe/średnie/duże" in stripped
        or "niepotrzebne skreślić" in stripped.lower()
    )


def _scan_tables(doc: Document, label_map: dict[str, str]) -> list[FieldResult]:
    results = []
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            cells = row.cells
            if len(cells) < 2:
                continue
            label_raw = cells[0].text.strip()
            value_raw = cells[1].text.strip()
            label_norm = _normalize_label(label_raw)
            if label_norm in label_map and _is_empty_value(value_raw):
                results.append(FieldResult(
                    location_id=f"T{t_idx}R{r_idx}",
                    location_type="table_cell",
                    label=label_raw,
                    original_value=value_raw,
                    filled_value=label_map[label_norm],
                    source=FieldSource.RULE,
                    confidence=Confidence.HIGH,
                ))
    return results


def _scan_paragraphs(doc: Document, label_map: dict[str, str]) -> list[FieldResult]:
    results = []
    for p_idx, para in enumerate(doc.paragraphs):
        text = para.text
        match = re.search(r'([\w\s\.\-]+?)\s*[:\.]?\s*([…\.]{3,})', text)
        if match:
            label_raw = re.sub(r'[:\*\.]+$', '', match.group(1).strip()).strip()
            label_norm = _normalize_label(label_raw)
            if label_norm in label_map:
                results.append(FieldResult(
                    location_id=f"P{p_idx}",
                    location_type="paragraph",
                    label=label_raw,
                    original_value=match.group(2),
                    filled_value=label_map[label_norm],
                    source=FieldSource.RULE,
                    confidence=Confidence.HIGH,
                ))
    return results


def analyze_rules(docx_path: Path, profile: CompanyProfile) -> list[FieldResult]:
    """Analyze DOCX and return fields that can be filled by rules."""
    doc = Document(str(docx_path))
    label_map = _build_label_map(profile)
    table_results = _scan_tables(doc, label_map)
    para_results = _scan_paragraphs(doc, label_map)
    return table_results + para_results
