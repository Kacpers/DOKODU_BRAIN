"""
POC: Regułowe uzupełnianie dokumentów przetargowych.
Testuje czy python-docx + regex wystarczy na prawdziwe formularze TenderScope.
"""
import os
import re
import copy
from docx import Document

# Przykładowy profil wykonawcy (hardcoded na potrzeby POC)
COMPANY_PROFILE = {
    "firma": "DOKODU Sp. z o.o.",
    "nazwa": "DOKODU Sp. z o.o.",
    "nazwa wykonawcy": "DOKODU Sp. z o.o.",
    "nazwa oferenta": "DOKODU Sp. z o.o.",
    "nip": "5882473305",
    "regon": "520149113",
    "krs": "0000925166",
    "adres": "ul. Kosynierów 76/22, 84-230 Rumia",
    "adres wykonawcy": "ul. Kosynierów 76/22, 84-230 Rumia",
    "adres do korespondencji": "ul. Kosynierów 76/22, 84-230 Rumia",
    "adres korespondencyjny": "ul. Kosynierów 76/22, 84-230 Rumia",
    "adres e-mail": "biuro@dokodu.it",
    "e-mail": "biuro@dokodu.it",
    "telefon": "+48 508 106 046",
    "numer telefonu": "+48 508 106 046",
    "osoba do kontaktu": "Kacper Sieradziński",
    "osoba do kontaktu\n(imię i nazwisko)": "Kacper Sieradziński",
    "miejscowość i data": "Rumia, 08.04.2026",
    "kapitał zakładowy": "5 000 PLN",
    "status przedsiębiorcy": "mikro",
}

# Normalizacja etykiet — lowercase, strip, usunięcie dwukropków
def normalize_label(text: str) -> str:
    t = text.strip().lower()
    t = re.sub(r'[:\*]+$', '', t).strip()
    t = re.sub(r'\s+', ' ', t)
    return t

def try_fill_table_cells(doc: Document) -> list[dict]:
    """Szuka tabel z etykietą w col0 i pustą komórką w col1."""
    fills = []
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            cells = row.cells
            if len(cells) < 2:
                continue
            label_raw = cells[0].text.strip()
            value_raw = cells[1].text.strip()
            label_norm = normalize_label(label_raw)

            # Komórka wartości pusta lub zawiera wielokropki/placeholder
            is_empty = (
                value_raw == ""
                or re.match(r'^[…\.]+$', value_raw)
                or value_raw == "mikro/małe/średnie/duże przedsiębiorstwo"
            )

            if label_norm in COMPANY_PROFILE and is_empty:
                new_value = COMPANY_PROFILE[label_norm]
                cells[1].paragraphs[0].text = new_value
                fills.append({
                    "method": "table_cell",
                    "table": t_idx,
                    "row": r_idx,
                    "label": label_raw,
                    "value": new_value,
                })
    return fills

def try_fill_placeholder_dots(doc: Document) -> list[dict]:
    """Szuka wielokropków w paragrafach (np. 'Łącznie brutto ……………')."""
    fills = []
    for p_idx, para in enumerate(doc.paragraphs):
        text = para.text
        # Szukaj wzorców typu "etykieta: ………" lub "etykieta ……………"
        match = re.search(r'([\w\s]+?)\s*[:]*\s*([…\.]{3,})', text)
        if match:
            label_norm = normalize_label(match.group(1))
            if label_norm in COMPANY_PROFILE:
                new_text = text.replace(match.group(2), COMPANY_PROFILE[label_norm])
                para.text = new_text
                fills.append({
                    "method": "paragraph_dots",
                    "paragraph": p_idx,
                    "label": match.group(1).strip(),
                    "value": COMPANY_PROFILE[label_norm],
                })
    return fills

def process_document(input_path: str, output_dir: str) -> dict:
    """Przetwarza jeden dokument, zwraca raport."""
    doc = Document(input_path)
    fname = os.path.basename(input_path)

    fills_table = try_fill_table_cells(doc)
    fills_dots = try_fill_placeholder_dots(doc)

    all_fills = fills_table + fills_dots

    if all_fills:
        out_path = os.path.join(output_dir, f"FILLED_{fname}")
        doc.save(out_path)

    return {
        "file": fname,
        "total_fills": len(all_fills),
        "table_fills": len(fills_table),
        "dot_fills": len(fills_dots),
        "details": all_fills,
    }

def main():
    input_dir = "/mnt/d/DOWNLOAD/reprobaowycenwykonaniemoduuuzupenianiadanychw"
    output_dir = "/tmp/tender_poc_output"
    os.makedirs(output_dir, exist_ok=True)

    results = []
    for fname in sorted(os.listdir(input_dir)):
        if not fname.endswith('.docx'):
            continue
        fpath = os.path.join(input_dir, fname)
        try:
            result = process_document(fpath, output_dir)
            results.append(result)
        except Exception as e:
            results.append({"file": fname, "error": str(e)})

    # Raport
    print("=" * 70)
    print("POC REGUŁOWY — RAPORT")
    print("=" * 70)
    total_fills = 0
    total_files_with_fills = 0
    for r in results:
        if "error" in r:
            print(f"\n❌ {r['file']}: {r['error']}")
            continue
        total_fills += r["total_fills"]
        if r["total_fills"] > 0:
            total_files_with_fills += 1
        print(f"\n📄 {r['file']}")
        print(f"   Wypełnione pola: {r['total_fills']} (tabele: {r['table_fills']}, wielokropki: {r['dot_fills']})")
        for d in r["details"]:
            print(f"   ✅ [{d['method']}] {d['label']} → {d['value']}")

    print(f"\n{'=' * 70}")
    print(f"PODSUMOWANIE: {total_fills} pól wypełnionych w {total_files_with_fills}/{len(results)} dokumentach")
    print(f"Pliki wyjściowe: {output_dir}")
    print("=" * 70)

if __name__ == "__main__":
    main()
